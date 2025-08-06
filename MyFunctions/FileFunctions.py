# Apresentação:
"""
Modulo de funções referentes à manipulação de arquivos.

A princípios vamos manter na manipulação de arquivos de texto,
fechando o escopo em md (markdown), pdf e docx (word).
"""

# Dependências
from pathlib import Path
from typing import Literal
import html
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm

def add_hyperlink(paragraph, url: str, label: str):
    """
    Adiciona um hyperlink em `paragraph`, com texto `label` apontando para `url`.
    Baseado em exemplo do GitHub python‑docx.:contentReference[oaicite:0]{index=0}
    """
    run = paragraph._p.add_hyperlink = None  # hack placeholder
    part = paragraph.part
    rid = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    link = OxmlElement('w:hyperlink')
    link.set(qn('r:id'), rid)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    run = paragraph.add_run(label)
    run._r.getparent().append(link)
    run._r.getparent().remove(run._r)

def save_to_file_logic(
    text: str,
    filename: str,
    fmt: str,
    dir_path: str = ""
) -> str:
    """
    Salva `text` (Markdown ou texto cru) em:
      • .md → grava como .md
      • .pdf → gera PDF com ReportLab + quebra automática (Paragraph)
      • .docx → gera Word (.docx) com python-docx, criando cabeçalhos e bullet simples
    """
    dir_ = Path(dir_path)
    dir_.mkdir(parents=True, exist_ok=True)
    full = dir_ / f"{filename}.{fmt}"

    if fmt == "md":
        full.write_text(text, encoding="utf8")
        return str(full)

    if fmt == "pdf":
        doc = SimpleDocTemplate(str(full), pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        normal = styles["Normal"]
        story = []
        for par in text.split("\n\n"):
            html_safe = html.escape(par).replace("\n", "<br />")
            story.append(Paragraph(html_safe, normal))
            story.append(Spacer(1, 0.3*cm))
        doc.build(story)
        return str(full)

    if fmt == "docx":
        doc = Document()
        for line in text.splitlines():
            if line.startswith("# "):
                doc.add_heading(line.removeprefix("# "), level=1)
            elif line.startswith("- "):
                doc.add_paragraph(line.removeprefix("- "), style="List Bullet")
            elif "[" in line and "](" in line and line.endswith(")"):
                label, url = line[1:].split("](")
                url = url[:-1]
                p = doc.add_paragraph()
                add_hyperlink(p, url, label)
            else:
                doc.add_paragraph(line)
        doc.save(str(full))
        return str(full)

    raise ValueError("Use fmt em 'md', 'pdf' ou 'docx'")